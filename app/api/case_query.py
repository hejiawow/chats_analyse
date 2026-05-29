# -*- coding: utf-8 -*-
"""优秀案例提取结果查询 API + 优秀成交案例 HTML 导出"""
import io
from datetime import datetime
from fastapi import APIRouter, Query, HTTPException, Depends
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy import select, func
from sqlalchemy.orm import Session

from app.models.database import sync_engine, async_session
from app.models.result import CaseExtractionResult, SalesJourneyResult
from app.models.script_lib import CaseScriptLibrary
from app.services.hujing_api import get_chat_records
from app.services.html_exporter import generate_html_report, DEFAULT_EXPORT_PATH
from app.services.dependencies import require_permission, get_current_user
from app.services.embedding import get_embeddings

router = APIRouter()


def _apply_user_filter(stmt, user_id: str | None, current_user: dict):
    """非 admin 用户自动过滤到当前 user_id"""
    if "admin:all" not in current_user.get("permissions", []):
        return stmt.where(CaseExtractionResult.user_id == str(current_user["user_id"]))
    if user_id:
        return stmt.where(CaseExtractionResult.user_id == user_id)
    return stmt


def _apply_user_filter_count(count_stmt, user_id: str | None, current_user: dict):
    if "admin:all" not in current_user.get("permissions", []):
        return count_stmt.where(CaseExtractionResult.user_id == str(current_user["user_id"]))
    if user_id:
        return count_stmt.where(CaseExtractionResult.user_id == user_id)
    return count_stmt


@router.get("/cases")
async def query_case_results(
    user_id: str | None = Query(None, description="销售ID"),
    friend_id: int | None = Query(None, description="好友ID"),
    script_type: str | None = Query(None, description="话术类型：销售话术 / 唤醒话术"),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    current_user: dict = Depends(require_permission("read:cases")),
):
    """查询优秀话术提取结果，支持按话术类型筛选。非 admin 只能查看自己的数据。"""
    async with async_session() as session:
        stmt = select(CaseExtractionResult)
        stmt = _apply_user_filter(stmt, user_id, current_user)
        if friend_id is not None:
            stmt = stmt.where(CaseExtractionResult.friend_id == friend_id)
        if script_type:
            stmt = stmt.where(CaseExtractionResult.script_type == script_type)

        stmt = stmt.order_by(CaseExtractionResult.created_at.desc())

        # 总数
        count_stmt = select(func.count()).select_from(CaseExtractionResult)
        count_stmt = _apply_user_filter_count(count_stmt, user_id, current_user)
        if friend_id is not None:
            count_stmt = count_stmt.where(CaseExtractionResult.friend_id == friend_id)
        if script_type:
            count_stmt = count_stmt.where(CaseExtractionResult.script_type == script_type)

        total = await session.scalar(count_stmt)

        # 分页数据
        stmt = stmt.offset((page - 1) * page_size).limit(page_size)
        result = await session.execute(stmt)
        records = result.scalars().all()

        return {
            "total": total,
            "page": page,
            "page_size": page_size,
            "data": [r.to_dict() for r in records],
        }


@router.get("/cases/{result_id}")
async def get_case_detail(
    result_id: int,
    current_user: dict = Depends(require_permission("read:cases")),
):
    """获取单条优秀案例提取详情"""
    async with async_session() as session:
        stmt = select(CaseExtractionResult).where(CaseExtractionResult.id == result_id)
        result = await session.execute(stmt)
        record = result.scalar_one_or_none()

        if not record:
            return {"error": "记录不存在"}

        # 非 admin 用户只能查看自己的数据
        if "admin:all" not in current_user.get("permissions", []):
            if record.user_id != str(current_user["user_id"]):
                return {"error": "无权查看此记录"}

        return record.to_dict()


@router.get("/cases/export/word")
async def export_cases_word(
    user_id: str | None = Query(None, description="销售ID"),
    script_type: str | None = Query(None, description="话术类型：销售话术 / 唤醒话术"),
    current_user: dict = Depends(require_permission("read:cases")),
):
    """导出优秀话术为 Word 文档"""
    async with async_session() as session:
        stmt = select(CaseExtractionResult).where(CaseExtractionResult.status == "success")
        # 非 admin 只能导出自己的数据
        if "admin:all" not in current_user.get("permissions", []):
            stmt = stmt.where(CaseExtractionResult.user_id == str(current_user["user_id"]))
        elif user_id:
            stmt = stmt.where(CaseExtractionResult.user_id == user_id)
        if script_type:
            stmt = stmt.where(CaseExtractionResult.script_type == script_type)
        stmt = stmt.order_by(CaseExtractionResult.created_at.desc())

        result = await session.execute(stmt)
        records = result.scalars().all()

    # Build Word document
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(4)

    # Title
    title = doc.add_heading("优秀话术提取报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "微软雅黑"
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Summary info
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = summary.add_run(f"生成时间：{now_str}")
    run1.font.size = Pt(9)
    run1.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run1.font.name = "微软雅黑"

    run2 = summary.add_run(f"\n案例总数：{len(records)}")
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run2.font.name = "微软雅黑"

    # Horizontal line
    doc.add_paragraph("_" * 70)

    for idx, record in enumerate(records, 1):
        # 话术标题（根据类型区分）
        if record.script_type == "唤醒话术":
            title_text = record.trigger_customer_state or "唤醒话术"
        else:
            title_text = record.customer_question or record.applicable_scenario or "销售话术"
        heading = doc.add_heading(f"话术 {idx}：{title_text[:50]}", level=2)
        for run in heading.runs:
            run.font.name = "微软雅黑"
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

        # 话术类型标签
        type_p = doc.add_paragraph()
        type_run = type_p.add_run(f"类型：{record.script_type or '未分类'}")
        type_run.bold = True
        type_run.font.size = Pt(10)
        type_run.font.name = "微软雅黑"
        type_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

        if record.script_type == "唤醒话术":
            # 唤醒话术字段展示
            display_fields = [
                ("触发客户状态", record.trigger_customer_state or "-"),
                ("销冠唤醒话术原文", record.wake_script or "-"),
                ("适用场景", record.applicable_scenario or "-"),
                ("话术核心目标", record.script_objective or "-"),
                ("适配人群", record.target_audience or "-"),
            ]
        else:
            # 销售话术字段展示
            display_fields = [
                ("客户问题", record.customer_question or "-"),
                ("销冠回答", record.sales_answer or "-"),
                ("适用场景", record.applicable_scenario or "-"),
                ("客户意图", record.customer_intent or "-"),
            ]

        for label, value in display_fields:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            label_run = p.add_run(f"{label}：")
            label_run.bold = True
            label_run.font.size = Pt(10.5)
            label_run.font.name = "微软雅黑"
            label_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            value_run = p.add_run(value)
            value_run.font.size = Pt(10.5)
            value_run.font.name = "微软雅黑"

        # 共享字段展示
        doc.add_paragraph("")
        shared_heading = doc.add_paragraph()
        shared_run = shared_heading.add_run("设计分析")
        shared_run.bold = True
        shared_run.font.size = Pt(11)
        shared_run.font.name = "微软雅黑"
        shared_run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

        shared_fields = [
            ("客户画像", record.customer_profile),
            ("核心设计逻辑", record.core_design_logic),
            ("话术关键技巧", record.key_techniques),
            ("标签", record.tags),
            ("业务科目", record.business_subject),
            ("合规风险", record.compliance_risk),
            ("反例避坑", record.pitfall_avoid),
        ]
        for label, value in shared_fields:
            if value:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(2)
                label_run = p.add_run(f"{label}：")
                label_run.bold = True
                label_run.font.size = Pt(10)
                label_run.font.name = "微软雅黑"
                label_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
                value_run = p.add_run(value)
                value_run.font.size = Pt(10)
                value_run.font.name = "微软雅黑"
                value_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        # Separator
        doc.add_paragraph("_" * 50)

    # Save to bytes buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"优秀话术_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename.encode('utf-8').decode('latin-1')}"},
    )


@router.get("/cases/export/html")
async def export_cases_html(
    user_id: str = Query(..., description="销售ID"),
    friend_id: int = Query(..., description="好友ID"),
    export_path: str = Query(DEFAULT_EXPORT_PATH, description="导出路径"),
    current_user: dict = Depends(require_permission("read:cases")),
):
    """
    生成优秀成交案例 HTML 复盘报告并写入本地文件。
    默认固定路径：C:\\Users\\23824\\Desktop\\玉娆高级职称王者班.html
    """
    # 1. 获取聊天记录
    chat_records = get_chat_records(user_id, friend_id)
    if not chat_records:
        raise HTTPException(status_code=404, detail="该好友无聊天记录")

    # 2. 获取最新一次案例提取分析结果
    async with async_session() as session:
        stmt = (
            select(SalesJourneyResult)
            .where(SalesJourneyResult.user_id == user_id)
            .where(SalesJourneyResult.friend_id == friend_id)
            .where(SalesJourneyResult.status == "success")
            .order_by(SalesJourneyResult.created_at.desc())
            .limit(1)
        )
        result = await session.execute(stmt)
        journey_record = result.scalar_one_or_none()

    if not journey_record or not journey_record.analysis_result:
        raise HTTPException(
            status_code=404,
            detail="未找到分析结果，请先触发分析（POST /api/trigger）",
        )

    analysis_result = journey_record.analysis_result
    friend_nick = journey_record.friend_nick or ""

    # 3. 生成 HTML 并写入本地文件
    file_path = generate_html_report(
        analysis_result=analysis_result,
        chat_records=chat_records,
        friend_nick=friend_nick,
        export_path=export_path,
    )

    return {
        "status": "success",
        "message": f"HTML 报告已生成并导出",
        "file_path": file_path,
        "friend_nick": friend_nick,
    }


def _build_document_text(case_dict: dict) -> str:
    """将话术字段拼接为 embedding 文本"""
    parts = []
    script_type = case_dict.get("script_type", "")

    if case_dict.get("applicable_scenario"):
        parts.append(f"场景：{case_dict['applicable_scenario']}")
    if script_type == "销售话术" and case_dict.get("customer_question"):
        parts.append(f"客户问题：{case_dict['customer_question']}")
    if script_type == "唤醒话术" and case_dict.get("trigger_customer_state"):
        parts.append(f"触发状态：{case_dict['trigger_customer_state']}")
    if case_dict.get("target_audience"):
        parts.append(f"适配人群：{case_dict['target_audience']}")
    if case_dict.get("customer_profile"):
        parts.append(f"客户画像：{case_dict['customer_profile']}")
    if case_dict.get("customer_intent"):
        parts.append(f"沟通阶段：{case_dict['customer_intent']}")

    quote = case_dict.get("sales_answer") or case_dict.get("wake_script", "")
    if quote:
        parts.append(f"销售话术：{quote}")
    if case_dict.get("core_design_logic"):
        parts.append(f"设计逻辑：{case_dict['core_design_logic']}")
    if case_dict.get("key_techniques"):
        parts.append(f"关键技巧：{case_dict['key_techniques']}")
    if case_dict.get("pitfall_avoid"):
        parts.append(f"反例避坑：{case_dict['pitfall_avoid']}")
    if case_dict.get("tags"):
        parts.append(f"标签：{case_dict['tags']}")
    if case_dict.get("business_subject"):
        parts.append(f"业务科目：{case_dict['business_subject']}")

    return "\n".join(parts)


@router.post("/cases/{case_id}/add-to-rag")
async def add_case_to_rag(
    case_id: int,
    current_user: dict = Depends(require_permission("write:scriptlib")),
):
    """将单条提取的话术案例存入话术库（生成 embedding 向量）"""
    # 1. 查询原始话术
    async with async_session() as session:
        stmt = select(CaseExtractionResult).where(CaseExtractionResult.id == case_id)
        result = await session.execute(stmt)
        case = result.scalar_one_or_none()

        if not case:
            raise HTTPException(status_code=404, detail="话术记录不存在")

        quote = case.sales_answer or case.wake_script
        if not quote or not quote.strip():
            raise HTTPException(status_code=400, detail="该记录无有效话术内容")

    # 2. 生成 embedding
    embeddings = get_embeddings()
    doc_text = _build_document_text(case.to_dict())

    if not doc_text.strip():
        raise HTTPException(status_code=400, detail="话术内容为空，无法生成 embedding")

    try:
        vec = embeddings.embed_documents([doc_text])[0]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"生成 embedding 失败：{str(e)}")

    # 3. 写入话术库
    with Session(sync_engine) as sync_session:
        # 检查是否已入库
        exist_count = sync_session.scalar(
            select(func.count()).where(CaseScriptLibrary.source_case_id == case_id)
        )
        if exist_count > 0:
            raise HTTPException(status_code=409, detail="该话术已存入库中")

        record = CaseScriptLibrary(
            source_case_id=case.id,
            user_id=case.user_id,
            user_wx_id=case.user_wx_id,
            friend_id=case.friend_id,
            friend_wx_id=case.friend_wx_id,
            friend_nick=case.friend_nick,
            scenario_type=case.applicable_scenario or "",
            customer_type=case.target_audience or "",
            communication_stage=case.customer_intent or "",
            sales_quote=quote,
            comprehensive_review=case.core_design_logic or "",
            customer_profile=case.customer_profile or "",
            document_text=doc_text,
            embedding=vec,
        )
        sync_session.add(record)
        sync_session.commit()

        return {
            "status": "success",
            "message": "话术已存入库",
            "rag_id": record.id,
            "script_type": case.script_type,
            "scenario": case.applicable_scenario,
        }


@router.post("/cases/batch-add-to-rag")
async def batch_add_cases_to_rag(
    user_id: str | None = Query(None),
    friend_id: int | None = Query(None),
    script_type: str | None = Query(None),
    current_user: dict = Depends(require_permission("write:scriptlib")),
):
    """批量将提取的话术案例存入话术库"""
    async with async_session() as session:
        stmt = select(CaseExtractionResult).where(CaseExtractionResult.status == "success")
        # 非 admin 只能操作自己的数据
        if "admin:all" not in current_user.get("permissions", []):
            stmt = stmt.where(CaseExtractionResult.user_id == str(current_user["user_id"]))
        elif user_id:
            stmt = stmt.where(CaseExtractionResult.user_id == user_id)
        if friend_id is not None:
            stmt = stmt.where(CaseExtractionResult.friend_id == friend_id)
        if script_type:
            stmt = stmt.where(CaseExtractionResult.script_type == script_type)

        result = await session.execute(stmt)
        cases = result.scalars().all()

    embeddings = get_embeddings()
    success_count = 0
    skip_count = 0
    fail_count = 0
    errors = []

    with Session(sync_engine) as sync_session:
        # 获取已入库的 source_case_id
        exist_ids = set(
            sync_session.execute(
                select(CaseScriptLibrary.source_case_id)
            ).scalars().all()
        )

        for case in cases:
            if case.id in exist_ids:
                skip_count += 1
                continue

            quote = case.sales_answer or case.wake_script
            if not quote or not quote.strip():
                skip_count += 1
                continue

            doc_text = _build_document_text(case.to_dict())
            if not doc_text.strip():
                fail_count += 1
                continue

            try:
                vec = embeddings.embed_documents([doc_text])[0]
                record = CaseScriptLibrary(
                    source_case_id=case.id,
                    user_id=case.user_id,
                    user_wx_id=case.user_wx_id,
                    friend_id=case.friend_id,
                    friend_wx_id=case.friend_wx_id,
                    friend_nick=case.friend_nick,
                    scenario_type=case.applicable_scenario or "",
                    customer_type=case.target_audience or "",
                    communication_stage=case.customer_intent or "",
                    sales_quote=quote,
                    comprehensive_review=case.core_design_logic or "",
                    customer_profile=case.customer_profile or "",
                    document_text=doc_text,
                    embedding=vec,
                )
                sync_session.add(record)
                sync_session.commit()
                success_count += 1
            except Exception as e:
                sync_session.rollback()
                fail_count += 1
                errors.append(f"ID {case.id}: {str(e)}")

    return {
        "status": "success",
        "added": success_count,
        "skipped": skip_count,
        "failed": fail_count,
        "errors": errors[:10],  # 最多返回10条错误
    }

